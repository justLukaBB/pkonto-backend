#!/usr/bin/env python3
"""
Script to fix checkboxes in Word template
Removes check from "geeignete Person" checkbox
"""

from docx import Document
import sys

def fix_checkboxes(docx_path):
    """Remove checkbox from 'geeignete Person'"""

    # Load document
    doc = Document(docx_path)

    changes_made = 0

    # Iterate through all paragraphs
    for para in doc.paragraphs:
        text = para.text

        # Check if this paragraph contains the checkbox text
        if 'geeignete Person' in text or 'geeignete Stelle' in text:
            print(f"Found: {text}")

            # Iterate through runs to find and replace checkbox symbols
            for run in para.runs:
                # Replace filled checkbox with empty checkbox for "geeignete Person"
                if '☑' in run.text and 'geeignete Person' in para.text:
                    run.text = run.text.replace('☑', '☐')
                    print(f"  ✓ Changed ☑ to ☐ in: {run.text}")
                    changes_made += 1

                # Also check for other checkbox symbols
                if '✓' in run.text and 'geeignete Person' in para.text:
                    run.text = run.text.replace('✓', '☐')
                    print(f"  ✓ Changed ✓ to ☐ in: {run.text}")
                    changes_made += 1

                if '✔' in run.text and 'geeignete Person' in para.text:
                    run.text = run.text.replace('✔', '☐')
                    print(f"  ✓ Changed ✔ to ☐ in: {run.text}")
                    changes_made += 1

    # Check tables as well
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text

                    if 'geeignete Person' in text or 'geeignete Stelle' in text:
                        print(f"Found in table: {text}")

                        for run in para.runs:
                            if '☑' in run.text and 'geeignete Person' in text:
                                run.text = run.text.replace('☑', '☐')
                                print(f"  ✓ Changed ☑ to ☐ in table: {run.text}")
                                changes_made += 1

                            if '✓' in run.text and 'geeignete Person' in text:
                                run.text = run.text.replace('✓', '☐')
                                print(f"  ✓ Changed ✓ to ☐ in table: {run.text}")
                                changes_made += 1

                            if '✔' in run.text and 'geeignete Person' in text:
                                run.text = run.text.replace('✔', '☐')
                                print(f"  ✓ Changed ✔ to ☐ in table: {run.text}")
                                changes_made += 1

    if changes_made > 0:
        # Save modified document
        backup_path = docx_path.replace('.docx', '.backup.docx')
        print(f"\n📦 Creating backup: {backup_path}")
        doc.save(backup_path)

        print(f"💾 Saving changes to: {docx_path}")
        doc.save(docx_path)
        print(f"\n✅ Successfully made {changes_made} change(s)!")
    else:
        print("\n⚠️  No checkboxes found to change. The document may use Word's native checkbox controls.")
        print("   Please open the document in Word and manually uncheck the 'geeignete Person' box.")

    return changes_made

if __name__ == '__main__':
    template_path = 'src/templates/certificate-template.docx'

    print("🔧 Fixing checkboxes in Word template...")
    print(f"📄 File: {template_path}\n")

    try:
        changes = fix_checkboxes(template_path)
        if changes == 0:
            print("\nℹ️  The document likely uses Word's native checkbox content controls.")
            print("   You'll need to open it in Word/LibreOffice and manually uncheck the box.")
    except Exception as e:
        print(f"\n❌ Error: {e}")
        print("\nMake sure python-docx is installed: pip install python-docx")
        sys.exit(1)
